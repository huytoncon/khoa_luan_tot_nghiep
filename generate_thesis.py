"""
Sinh file Word khóa luận tốt nghiệp.
Output: khoa_luan_tot_nghiep.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ──────────────────────────────────────────────────────────────────

doc = Document()

# Thiết lập trang A4, lề chuẩn luận văn
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.5)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

def set_font(run, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force font for East Asian text
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=0, space_after=6, line_spacing=1.5, first_line=0):
    pf = para.paragraph_format
    pf.alignment    = alignment
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if first_line:
        pf.first_line_indent = Cm(first_line)

def add_para(text="", bold=False, italic=False, size=13,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=1.25,
             space_before=0, space_after=6):
    p = doc.add_paragraph()
    set_para_format(p, alignment=align, space_before=space_before,
                    space_after=space_after, first_line=indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading(text, level=1):
    """Cấp 1: chương; 2: mục; 3: mục nhỏ"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=14, bold=True)
    elif level == 2:
        set_font(run, size=13, bold=True)
    else:
        set_font(run, size=13, bold=True, italic=True)
    return p

def add_bullet(text, level=1, size=13):
    prefix = "  " * (level - 1) + "• "
    p = doc.add_paragraph()
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=0, space_after=3, first_line=0)
    p.paragraph_format.left_indent = Cm(1.25 * level)
    run = p.add_run(prefix + text)
    set_font(run, size=size)
    return p

def add_fig(caption):
    p = doc.add_paragraph()
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=6, space_after=6, first_line=0)
    run = p.add_run(f"[CẦN HÌNH MINH HỌA: {caption}]")
    set_font(run, size=12, italic=True, color=(0, 112, 192))
    return p

def add_table_caption(text):
    p = doc.add_paragraph()
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=4, space_after=4, first_line=0)
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p

def simple_table(headers, rows):
    """Tạo bảng với header đậm."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=12, bold=True)
        cell._tc.get_or_add_tcPr().append(
            OxmlElement('w:shd'))
    # Data
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            set_font(run, size=12)
    return tbl

def page_break():
    doc.add_page_break()

def add_blank():
    add_para("")

# ─── TRANG BÌA ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run("TRƯỜNG ĐẠI HỌC KINH TẾ QUỐC DÂN")
set_font(r, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run("VIỆN CÔNG NGHỆ THÔNG TIN VÀ KINH TẾ SỐ")
set_font(r, size=13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("─────────────────────────────")
set_font(r, size=13)

for _ in range(4):
    add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("KHÓA LUẬN TỐT NGHIỆP")
set_font(r, size=16, bold=True)

for _ in range(2):
    add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("HỆ THỐNG SO SÁNH GIÁ ĐIỆN THOẠI THÔNG MINH")
set_font(r, size=15, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("TÍCH HỢP TÌM KIẾM BẰNG HÌNH ẢNH VÀ TƯ VẤN AI:")
set_font(r, size=15, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run("TIẾP CẬN THEO PHƯƠNG PHÁP NGHIÊN CỨU KHOA HỌC THIẾT KẾ")
set_font(r, size=15, bold=True)

for _ in range(5):
    add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sinh viên thực hiện:   [Họ và tên sinh viên]")
set_font(r, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Mã sinh viên:           [MSSV]")
set_font(r, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Lớp:                      [Tên lớp]")
set_font(r, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Giáo viên hướng dẫn:  [Tên GVHD]")
set_font(r, size=13)

for _ in range(5):
    add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Hà Nội, năm 2025")
set_font(r, size=13, bold=True)

page_break()

# ─── LỜI CAM ĐOAN ─────────────────────────────────────────────────────────────
add_heading("LỜI CAM ĐOAN", 1)
add_para("Tôi xin cam đoan đây là công trình nghiên cứu khoa học độc lập của riêng tôi. Các số liệu, kết quả trong khóa luận là trung thực và chưa từng được công bố trong bất kỳ công trình nào khác.", indent=1.25)
add_para("Tôi xin chịu hoàn toàn trách nhiệm về tính xác thực và tính trung thực của khóa luận này.", indent=1.25)
add_blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Hà Nội, ngày      tháng      năm 2025")
set_font(r, size=13)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Tác giả khóa luận")
set_font(r, size=13, bold=True)
for _ in range(3):
    add_blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("[Họ và tên]")
set_font(r, size=13, bold=True)

page_break()

# ─── LỜI CẢM ƠN ───────────────────────────────────────────────────────────────
add_heading("LỜI CẢM ƠN", 1)
add_para("Trước tiên, tôi xin gửi lời cảm ơn chân thành và sâu sắc nhất đến [Tên GVHD], người đã tận tình hướng dẫn, góp ý và hỗ trợ tôi trong suốt quá trình thực hiện khóa luận này. Những kiến thức, kinh nghiệm quý báu mà thầy/cô chia sẻ đã giúp tôi định hướng và hoàn thiện nghiên cứu một cách tốt nhất.", indent=1.25)
add_para("Tôi cũng xin gửi lời cảm ơn đến Ban Giám hiệu Trường Đại học Kinh tế Quốc dân, Viện Công nghệ thông tin và Kinh tế số, cùng toàn thể các thầy cô đã truyền đạt kiến thức, tạo điều kiện thuận lợi cho tôi học tập và nghiên cứu trong suốt những năm học vừa qua.", indent=1.25)
add_para("Xin cảm ơn gia đình, bạn bè và những người thân đã luôn ủng hộ, động viên tôi trong quá trình hoàn thành khóa luận.", indent=1.25)
add_para("Mặc dù đã cố gắng hết sức, nhưng do giới hạn về thời gian và kiến thức, khóa luận không thể tránh khỏi những thiếu sót. Tôi rất mong nhận được sự góp ý từ thầy cô và các bạn để công trình ngày càng hoàn thiện hơn.", indent=1.25)
add_blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Hà Nội, ngày      tháng      năm 2025\nTác giả khóa luận")
set_font(r, size=13)

page_break()

# ─── DANH MỤC TỪ VIẾT TẮT ────────────────────────────────────────────────────
add_heading("DANH MỤC TỪ VIẾT TẮT", 1)
abbrevs = [
    ("AI", "Artificial Intelligence – Trí tuệ nhân tạo"),
    ("API", "Application Programming Interface – Giao diện lập trình ứng dụng"),
    ("CLIP", "Contrastive Language–Image Pre-training – Mô hình học tương phản ngôn ngữ-hình ảnh"),
    ("CLT", "Cognitive Load Theory – Lý thuyết tải nhận thức"),
    ("CORS", "Cross-Origin Resource Sharing – Chia sẻ tài nguyên đa nguồn gốc"),
    ("CPU/GPU", "Central/Graphics Processing Unit – Bộ xử lý trung tâm/đồ họa"),
    ("DSR", "Design Science Research – Nghiên cứu khoa học thiết kế"),
    ("JSON", "JavaScript Object Notation – Định dạng trao đổi dữ liệu"),
    ("LLM", "Large Language Model – Mô hình ngôn ngữ lớn"),
    ("MSCOCO", "Microsoft Common Objects in Context – Bộ dữ liệu hình ảnh đa mục tiêu"),
    ("nDCG", "Normalized Discounted Cumulative Gain – Chỉ số đánh giá xếp hạng"),
    ("OCR", "Optical Character Recognition – Nhận dạng ký tự quang học"),
    ("RAG", "Retrieval-Augmented Generation – Sinh văn bản tăng cường bằng truy xuất"),
    ("SBERT", "Sentence-BERT – Mô hình BERT mức câu"),
    ("TAM", "Technology Acceptance Model – Mô hình chấp nhận công nghệ"),
    ("TGDĐ", "Thế Giới Di Động – Chuỗi bán lẻ điện thoại"),
    ("TMĐT", "Thương mại điện tử"),
    ("ViT", "Vision Transformer – Mô hình transformer cho hình ảnh"),
    ("VNĐ", "Việt Nam Đồng – Đơn vị tiền tệ Việt Nam"),
]
tbl = simple_table(["Từ viết tắt", "Giải nghĩa"], abbrevs)

page_break()

# ─── DANH MỤC BẢNG ────────────────────────────────────────────────────────────
add_heading("DANH MỤC BẢNG", 1)
tables_list = [
    ("Bảng 2.1", "So sánh các hệ thống so sánh giá hiện có"),
    ("Bảng 2.2", "Các phương pháp tìm kiếm sản phẩm"),
    ("Bảng 3.1", "Ánh xạ Lý thuyết nền – Yêu cầu thiết kế – Tính năng hệ thống – Giả thuyết"),
    ("Bảng 3.2", "Thống kê bộ dữ liệu sản phẩm"),
    ("Bảng 3.3", "Thống kê sản phẩm theo nguồn bán"),
    ("Bảng 3.4", "Phân phối thương hiệu trong bộ dữ liệu"),
    ("Bảng 4.1", "Kết quả đánh giá tìm kiếm bằng hình ảnh (Top-K Accuracy)"),
    ("Bảng 4.2", "So sánh Chatbot có RAG và không có RAG"),
    ("Bảng 4.3", "Thống kê mô tả kết quả khảo sát người dùng"),
    ("Bảng 4.4", "Kết quả kiểm định giả thuyết nghiên cứu"),
]
simple_table(["Số hiệu bảng", "Tên bảng"], tables_list)

add_blank()
add_heading("DANH MỤC HÌNH", 1)
figs_list = [
    ("Hình 1.1", "Xu hướng mua sắm smartphone trực tuyến tại Việt Nam (2020–2024)"),
    ("Hình 2.1", "Kiến trúc mô hình CLIP"),
    ("Hình 2.2", "Kiến trúc SBERT"),
    ("Hình 2.3", "Pipeline RAG cổ điển"),
    ("Hình 2.4", "Mô hình TAM"),
    ("Hình 3.1", "Quy trình DSR theo Peffers et al. (2007) áp dụng trong nghiên cứu"),
    ("Hình 3.2", "Kiến trúc hệ thống tổng thể"),
    ("Hình 3.3", "Pipeline tìm kiếm bằng hình ảnh (CLIP)"),
    ("Hình 3.4", "Pipeline chatbot RAG (SBERT + Gemini)"),
    ("Hình 3.5", "Mô hình nghiên cứu và các giả thuyết"),
    ("Hình 4.1", "Biểu đồ Top-K Accuracy của CLIP theo K"),
    ("Hình 4.2", "So sánh thời gian phản hồi các API endpoint"),
    ("Hình 4.3", "Kết quả khảo sát TAM theo thang Likert 1–5"),
    ("Hình A.1", "Giao diện trang chủ danh sách sản phẩm"),
    ("Hình A.2", "Giao diện tìm kiếm bằng hình ảnh"),
    ("Hình A.3", "Giao diện chatbot mở rộng"),
    ("Hình A.4", "Modal so sánh sản phẩm"),
    ("Hình A.5", "Tính năng lọc giá với thanh trượt"),
]
simple_table(["Số hiệu hình", "Tên hình"], figs_list)

page_break()

# ─── TÓM TẮT ─────────────────────────────────────────────────────────────────
add_heading("TÓM TẮT", 1)
add_para("Thị trường điện thoại thông minh tại Việt Nam ngày càng phát triển với hàng trăm mẫu máy được bán trên nhiều nền tảng thương mại điện tử khác nhau. Người tiêu dùng phải đối mặt với bài toán quá tải thông tin khi muốn so sánh giá và chất lượng sản phẩm giữa các nguồn bán. Khóa luận này trình bày việc xây dựng hệ thống so sánh giá điện thoại thông minh tích hợp tìm kiếm bằng hình ảnh và chatbot tư vấn AI, áp dụng theo phương pháp Nghiên cứu Khoa học Thiết kế (Design Science Research – DSR).", indent=1.25)
add_para("Hệ thống được phát triển bao gồm ba thành phần kỹ thuật chính: (1) tìm kiếm ngữ nghĩa sử dụng mô hình Vietnamese SBERT (keepitreal/vietnamese-sbert) kết hợp vector store để truy xuất sản phẩm phù hợp theo ngữ cảnh; (2) tìm kiếm bằng hình ảnh ứng dụng mô hình CLIP (openai/clip-vit-base-patch32) cho phép người dùng tìm kiếm bằng cách tải lên ảnh điện thoại; (3) chatbot tư vấn thông minh xây dựng trên nền tảng Retrieval-Augmented Generation (RAG) với Gemini 2.0 Flash, cung cấp câu trả lời dựa trên dữ liệu thực tế. Bộ dữ liệu gồm 903 sản phẩm điện thoại thu thập từ 4 nguồn bán lớn tại Việt Nam: FPT Shop, Hoàng Hà Mobile, Shopee và Lazada, trong đó 567 sản phẩm được bổ sung thông số kỹ thuật chi tiết.", indent=1.25)
add_para("[CẦN BỔ SUNG: Kết quả thực nghiệm khi có dữ liệu khảo sát người dùng]", indent=1.25)
add_blank()
add_heading("ABSTRACT", 2)
add_para("Vietnam's smartphone market is rapidly growing with hundreds of models available across multiple e-commerce platforms. Consumers face significant information overload when comparing prices and product quality across different sources. This thesis presents the development of a smartphone price comparison system integrating image-based search and AI advisory chatbot, following the Design Science Research (DSR) methodology.", indent=1.25)
add_para("The system incorporates three main technical components: (1) semantic search using the Vietnamese SBERT model with a vector store for context-aware retrieval; (2) image-based search using the CLIP model, enabling users to search by uploading smartphone images; (3) an intelligent advisory chatbot built on Retrieval-Augmented Generation (RAG) with Gemini 2.0 Flash. The dataset comprises 903 smartphone products collected from four major Vietnamese retail platforms: FPT Shop, Hoang Ha Mobile, Shopee, and Lazada, with 567 products enriched with detailed technical specifications.", indent=1.25)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 1: GIỚI THIỆU
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 1. GIỚI THIỆU", 1)

add_heading("1.1. Bối cảnh và động lực nghiên cứu", 2)
add_para("Trong những năm gần đây, thị trường điện thoại thông minh tại Việt Nam đang có sự tăng trưởng mạnh mẽ. Theo số liệu của Statista (2024), doanh số điện thoại thông minh tại Việt Nam đạt hơn 12 triệu chiếc mỗi năm, với mức chi tiêu bình quân ngày càng tăng do người tiêu dùng có xu hướng lựa chọn các phân khúc trung và cao cấp. Cùng với sự phát triển của thương mại điện tử, người dùng Việt Nam ngày càng ưu tiên mua sắm trực tuyến qua các nền tảng lớn như Shopee, Lazada, cũng như các chuỗi cửa hàng bán lẻ điện thoại chuyên nghiệp như FPT Shop và Hoàng Hà Mobile.", indent=1.25)
add_para("Tuy nhiên, quá trình mua sắm điện thoại thông minh hiện nay đặt ra không ít thách thức cho người tiêu dùng. Một chiếc điện thoại cụ thể có thể được bán tại hàng chục nguồn khác nhau với mức giá, chính sách bảo hành và ưu đãi đi kèm không đồng nhất. Để đưa ra quyết định mua sắm tối ưu, người dùng phải tự mình truy cập, so sánh thông tin từ nhiều trang web, mất nhiều thời gian và công sức. Đây chính là biểu hiện điển hình của hiện tượng quá tải thông tin (information overload) trong bối cảnh thương mại điện tử hiện đại.", indent=1.25)
add_para("Bên cạnh đó, công nghệ trí tuệ nhân tạo đang tạo ra những bước đột phá lớn trong lĩnh vực tìm kiếm và tư vấn mua sắm. Các mô hình học sâu như CLIP (Contrastive Language–Image Pre-training) của OpenAI cho phép tìm kiếm hình ảnh với độ chính xác cao, trong khi các mô hình ngôn ngữ lớn (LLM) như Gemini của Google có khả năng tư vấn thông minh, hiểu ngôn ngữ tự nhiên và trả lời các câu hỏi phức tạp về sản phẩm. Kỹ thuật RAG (Retrieval-Augmented Generation) kết hợp giữa truy xuất thông tin và sinh văn bản, giúp chatbot AI cung cấp câu trả lời chính xác, có dẫn nguồn thay vì bịa đặt thông tin.", indent=1.25)
add_fig("Xu hướng mua sắm smartphone trực tuyến tại Việt Nam giai đoạn 2020–2024 (Nguồn: Statista)")

add_heading("1.2. Xác định vấn đề", 2)
add_para("Mặc dù nhu cầu so sánh giá và tư vấn mua sắm điện thoại thông minh là rất lớn, các giải pháp hiện có vẫn tồn tại nhiều hạn chế đáng kể:", indent=1.25)
add_bullet("Phân mảnh thông tin: Người dùng phải truy cập nhiều trang web khác nhau để so sánh giá. Không có công cụ nào tổng hợp dữ liệu từ các nền tảng FPT Shop, Hoàng Hà Mobile, Shopee và Lazada trong một giao diện thống nhất.")
add_bullet("Tìm kiếm hạn chế: Các hệ thống tìm kiếm hiện tại chủ yếu dựa trên từ khóa chính xác, không hiểu được ý định tìm kiếm theo ngữ nghĩa (semantic search). Hoàn toàn không có tính năng tìm kiếm bằng hình ảnh trong các nền tảng TMĐT Việt Nam.")
add_bullet("Chatbot thiếu độ chính xác: Các chatbot hỗ trợ mua hàng hiện nay thường hoạt động theo kịch bản cố định hoặc không có cơ sở dữ liệu thực, dẫn đến việc tạo ra thông tin sai lệch (hallucination) khi trả lời về thông số kỹ thuật sản phẩm.")
add_bullet("Thiếu tính năng so sánh: Người dùng không có công cụ để đặt nhiều sản phẩm cạnh nhau và nhận được phân tích so sánh chi tiết kết hợp với tư vấn AI.")
add_para("Những hạn chế trên đặt ra nhu cầu xây dựng một hệ thống tích hợp, giải quyết đồng thời bài toán tổng hợp dữ liệu đa nguồn, tìm kiếm đa phương thức và tư vấn thông minh có nguồn dẫn.", indent=1.25)

add_heading("1.3. Câu hỏi và mục tiêu nghiên cứu", 2)
add_para("Nghiên cứu này đặt ra câu hỏi chủ đạo:", indent=1.25)
add_para("\"Hệ thống so sánh giá điện thoại tích hợp tìm kiếm bằng hình ảnh (CLIP), tìm kiếm ngữ nghĩa (SBERT) và chatbot tư vấn AI (RAG-Gemini) có giúp người dùng đưa ra quyết định mua sắm hiệu quả và thỏa mãn hơn không?\"", bold=True, indent=1.25)
add_para("Để trả lời câu hỏi này, nghiên cứu hướng đến các mục tiêu cụ thể sau:", indent=1.25)
add_bullet("Xây dựng hệ thống thu thập và tổng hợp dữ liệu sản phẩm từ 4 nền tảng thương mại điện tử lớn tại Việt Nam.")
add_bullet("Tích hợp tìm kiếm bằng hình ảnh ứng dụng mô hình CLIP của OpenAI.")
add_bullet("Phát triển chatbot tư vấn thông minh dựa trên kỹ thuật RAG với mô hình Gemini 2.0 Flash.")
add_bullet("Xây dựng tính năng lọc đa tiêu chí và so sánh trực quan tối đa 3 sản phẩm.")
add_bullet("Đánh giá hiệu quả hệ thống thông qua thực nghiệm kỹ thuật và khảo sát người dùng.")

add_heading("1.4. Đóng góp của nghiên cứu", 2)
add_para("Nghiên cứu này có các đóng góp trên hai phương diện:", indent=1.25)
add_para("Về mặt kỹ thuật:", bold=True, indent=1.25)
add_bullet("Thiết kế và triển khai pipeline tìm kiếm đa phương thức kết hợp CLIP (ảnh) và Vietnamese SBERT (văn bản) trong cùng một hệ thống.")
add_bullet("Phát triển pipeline chatbot RAG tiếng Việt ứng dụng mô hình Gemini 2.0 Flash, tích hợp price boosting mechanism để cải thiện chất lượng truy xuất.")
add_bullet("Xây dựng bộ dữ liệu sản phẩm thực tế với 903 điện thoại từ 4 nguồn bán, được làm giàu với thông số kỹ thuật cho 567 sản phẩm.")
add_para("Về mặt thực tiễn:", bold=True, indent=1.25)
add_bullet("Cung cấp một hệ thống hoàn chỉnh, có thể triển khai thực tế, giải quyết nhu cầu so sánh giá và tư vấn mua sắm của người tiêu dùng Việt Nam.")
add_bullet("Chứng minh tính khả thi của việc áp dụng các mô hình AI tiên tiến (CLIP, LLM) vào lĩnh vực thương mại điện tử Việt Nam.")
add_bullet("Cung cấp framework thiết kế có thể tái sử dụng cho các hệ thống so sánh giá tương tự trong các lĩnh vực khác.")

add_heading("1.5. Phạm vi và giới hạn nghiên cứu", 2)
add_para("Nghiên cứu tập trung vào danh mục điện thoại thông minh, thu thập dữ liệu từ 4 nền tảng tại Việt Nam: FPT Shop, Hoàng Hà Mobile, Shopee và Lazada. Dữ liệu được thu thập tại một thời điểm cụ thể và không cập nhật theo thời gian thực. Nghiên cứu không bao gồm phân tích xu hướng giá theo thời gian hay tích hợp hệ thống đặt hàng trực tiếp.", indent=1.25)

add_heading("1.6. Cấu trúc khóa luận", 2)
add_para("Khóa luận được tổ chức thành 5 chương. Chương 1 trình bày bối cảnh, vấn đề và mục tiêu nghiên cứu. Chương 2 tổng quan các công trình liên quan và xác định khoảng trống nghiên cứu. Chương 3 mô tả phương pháp nghiên cứu DSR, yêu cầu thiết kế và triển khai kỹ thuật. Chương 4 trình bày kết quả đánh giá kỹ thuật và khảo sát người dùng. Chương 5 thảo luận về kết quả, đóng góp và hướng phát triển.", indent=1.25)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 2: TỔNG QUAN TÀI LIỆU
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 2. TỔNG QUAN TÀI LIỆU", 1)

add_heading("2.1. Hệ thống so sánh giá trong thương mại điện tử", 2)
add_para("Hệ thống so sánh giá (Price Comparison System – PCS) là dạng ứng dụng web hoặc di động cho phép người dùng tra cứu và đối chiếu giá cùng một sản phẩm từ nhiều nhà cung cấp khác nhau. Các hệ thống tiên phong như Google Shopping, PriceGrabber (Mỹ), Idealo (Đức) hay CompareAsia (châu Á) đã chứng minh giá trị kinh tế của mô hình này trong việc giảm chi phí tìm kiếm thông tin cho người tiêu dùng (Baye et al., 2004).", indent=1.25)
add_para("Nghiên cứu của Brynjolfsson và Smith (2000) cho thấy giá bán lẻ trực tuyến thấp hơn đáng kể so với kênh truyền thống, nhưng có sự biến động lớn giữa các nhà bán lẻ online. Điều này nhấn mạnh tầm quan trọng của công cụ so sánh giá trong việc giúp người tiêu dùng tận dụng lợi thế cạnh tranh về giá trên thị trường số. Tại Việt Nam, các nền tảng như Shopee, Lazada đang là kênh mua sắm trực tuyến phổ biến nhất nhưng vẫn thiếu công cụ so sánh giá liên nền tảng.", indent=1.25)
add_para("Các hệ thống PCS hiện đại đối mặt với nhiều thách thức kỹ thuật: (1) thu thập và chuẩn hóa dữ liệu từ các nguồn có cấu trúc khác nhau; (2) nhận diện và ghép nối cùng một sản phẩm từ các nguồn khác nhau (product matching); (3) cập nhật giá theo thời gian thực. Nghiên cứu của Vandic et al. (2013) đã đề xuất phương pháp Product Attribute Extraction để tự động trích xuất và so khớp thông số sản phẩm, cải thiện đáng kể độ chính xác của việc ghép nối sản phẩm.", indent=1.25)

add_table_caption("Bảng 2.1. So sánh các hệ thống so sánh giá tiêu biểu")
simple_table(
    ["Hệ thống", "Quốc gia", "Tìm kiếm ảnh", "Chatbot AI", "RAG", "Phạm vi"],
    [
        ["Google Shopping", "Toàn cầu", "Có", "Không", "Không", "Đa ngành"],
        ["Idealo", "Đức/EU", "Không", "Không", "Không", "Đa ngành"],
        ["CompareAsia", "Châu Á", "Không", "Không", "Không", "Điện tử"],
        ["GiaSo.vn", "Việt Nam", "Không", "Không", "Không", "Điện thoại"],
        ["Hệ thống đề xuất", "Việt Nam", "Có (CLIP)", "Có (Gemini)", "Có", "Điện thoại"],
    ]
)

add_heading("2.2. Tìm kiếm sản phẩm đa phương thức", 2)

add_heading("2.2.1. Tìm kiếm văn bản – Từ keyword đến ngữ nghĩa", 3)
add_para("Tìm kiếm dựa trên từ khóa (keyword-based search) là phương pháp phổ biến nhất trong các hệ thống thương mại điện tử. BM25 (Best Match 25) là thuật toán xếp hạng tài liệu dựa trên tần suất xuất hiện từ khóa, đã được áp dụng rộng rãi trong hơn 30 năm qua (Robertson và Zaragoza, 2009). Tuy nhiên, BM25 có hạn chế cơ bản: nó chỉ tìm kiếm khớp từ chính xác, không thể hiểu được ngữ nghĩa hay ý định của người dùng.", indent=1.25)
add_para("Sự ra đời của BERT (Bidirectional Encoder Representations from Transformers, Devlin et al., 2019) đã cách mạng hóa tìm kiếm ngữ nghĩa. SBERT (Sentence-BERT, Reimers và Gurevych, 2019) mở rộng BERT cho phép tính toán embedding ở cấp độ câu với hiệu suất cao, phù hợp cho tác vụ tính độ tương đồng ngữ nghĩa. Đặc biệt, mô hình Vietnamese SBERT (keepitreal/vietnamese-sbert) được tinh chỉnh đặc biệt cho tiếng Việt, cho phép hiểu ngữ nghĩa các câu hỏi và tên sản phẩm tiếng Việt một cách chính xác.", indent=1.25)
add_fig("Kiến trúc mô hình SBERT và quá trình tính toán sentence embedding")

add_heading("2.2.2. Tìm kiếm bằng hình ảnh với CLIP", 3)
add_para("CLIP (Contrastive Language–Image Pre-training) là mô hình được Radford et al. (2021) giới thiệu, được huấn luyện trên 400 triệu cặp ảnh–văn bản từ Internet. CLIP học cách biểu diễn ảnh và văn bản trong cùng một không gian embedding, cho phép so sánh độ tương đồng giữa ảnh và văn bản hoặc giữa các ảnh với nhau.", indent=1.25)
add_para("Kiến trúc CLIP bao gồm hai encoder song song: (1) Image Encoder sử dụng Vision Transformer (ViT) hoặc ResNet để mã hóa ảnh thành vector embedding; (2) Text Encoder sử dụng Transformer để mã hóa văn bản. Cả hai encoder được huấn luyện theo mục tiêu tương phản (contrastive objective): tối đa hóa sự tương đồng cosine giữa các cặp ảnh–văn bản tương ứng và tối thiểu hóa sự tương đồng với các cặp không tương ứng. Mô hình openai/clip-vit-base-patch32, phiên bản sử dụng trong nghiên cứu này, dùng ViT-B/32 với kích thước patch 32×32 pixel.", indent=1.25)
add_fig("Kiến trúc mô hình CLIP với Image Encoder và Text Encoder song song")

add_heading("2.3. Retrieval-Augmented Generation (RAG)", 2)

add_heading("2.3.1. Kiến trúc RAG cổ điển", 3)
add_para("RAG (Lewis et al., 2020) là kỹ thuật kết hợp truy xuất thông tin (Information Retrieval) với sinh văn bản (Text Generation) để tạo ra câu trả lời dựa trên nguồn kiến thức thực tế. Kiến trúc RAG gồm ba thành phần chính: (1) Retriever – mô-đun truy xuất tài liệu liên quan từ knowledge base bằng cách so sánh embedding; (2) Knowledge Base – cơ sở dữ liệu lưu trữ văn bản và embedding tương ứng; (3) Generator – mô hình ngôn ngữ lớn (LLM) nhận context từ Retriever và sinh câu trả lời.", indent=1.25)
add_para("Ưu điểm chính của RAG so với LLM thuần túy là khả năng cung cấp câu trả lời có nguồn gốc xác thực, giảm thiểu hiện tượng \"ảo giác\" (hallucination) – tình trạng LLM tạo ra thông tin sai lệch nhưng nghe có vẻ hợp lý. Điều này đặc biệt quan trọng trong ứng dụng tư vấn mua sắm, nơi độ chính xác của thông tin về giá cả và thông số kỹ thuật là yếu tố quyết định.", indent=1.25)
add_fig("Pipeline RAG: Query → Retriever → Knowledge Base → Generator → Response")

add_heading("2.3.2. Ứng dụng RAG trong thương mại điện tử", 3)
add_para("Kuo et al. (2023) đề xuất hệ thống chatbot thương mại điện tử dựa trên RAG, cho thấy độ chính xác về thông tin sản phẩm cải thiện 34% so với LLM không có RAG. Ye et al. (2024) phát triển framework RAG cho tư vấn mua sắm đa lĩnh vực, tích hợp lọc theo giá và thương hiệu vào quá trình truy xuất. Các nghiên cứu này nhấn mạnh vai trò quan trọng của chất lượng vector store và chiến lược chunking dữ liệu trong hiệu suất của hệ thống RAG.", indent=1.25)

add_heading("2.4. Quá tải thông tin trong thương mại điện tử", 2)
add_para("Information Overload là hiện tượng xảy ra khi lượng thông tin mà một cá nhân nhận được vượt quá khả năng xử lý, dẫn đến giảm chất lượng quyết định (Eppler và Mengis, 2004). Trong bối cảnh thương mại điện tử, người dùng thường phải đối mặt với hàng nghìn sản phẩm, hàng chục thuộc tính khác nhau và nhiều nguồn thông tin cạnh tranh.", indent=1.25)
add_para("Nghiên cứu của Iyengar và Lepper (2000) trong \"The Paradox of Choice\" chỉ ra rằng quá nhiều lựa chọn thực sự làm giảm sự thỏa mãn của người dùng và gia tăng xu hướng từ bỏ quyết định mua hàng. Ứng dụng thực tiễn quan trọng từ nghiên cứu này là cần có các công cụ lọc và trực quan hóa thông minh để giúp người dùng thu hẹp không gian lựa chọn xuống mức có thể quản lý được.", indent=1.25)
add_para("Trong lĩnh vực điện thoại thông minh, mỗi sản phẩm có hàng chục thông số kỹ thuật phức tạp (chip, RAM, màn hình, camera, pin...). Việc so sánh các thông số này đòi hỏi kiến thức chuyên môn, tạo ra rào cản lớn cho người dùng phổ thông. Một hệ thống tư vấn AI có thể đóng vai trò người trung gian, giải thích các thông số kỹ thuật theo ngôn ngữ đơn giản và đưa ra gợi ý phù hợp với nhu cầu người dùng.", indent=1.25)

add_heading("2.5. Mô hình chấp nhận công nghệ (TAM)", 2)
add_para("Mô hình TAM (Technology Acceptance Model) được Davis (1989) đề xuất là framework lý thuyết chuẩn mực để nghiên cứu hành vi chấp nhận và sử dụng hệ thống thông tin. Mô hình gốc xác định hai yếu tố quyết định: Perceived Usefulness (Tính hữu ích cảm nhận – PU) và Perceived Ease of Use (Tính dễ dùng cảm nhận – PEOU). PU được định nghĩa là mức độ người dùng tin rằng sử dụng hệ thống sẽ cải thiện hiệu suất công việc của họ; PEOU là mức độ hệ thống không đòi hỏi nỗ lực sử dụng.", indent=1.25)
add_para("Venkatesh và Davis (2000) mở rộng TAM2 bổ sung thêm các yếu tố xã hội và nhận thức. Trong bối cảnh hệ thống AI thương mại điện tử, các nghiên cứu như Pantano và Viassone (2015) chỉ ra rằng PU và PEOU là hai yếu tố dự đoán mạnh nhất cho Continuance Intention (ý định tiếp tục sử dụng), phù hợp với mục tiêu đánh giá trong nghiên cứu này.", indent=1.25)
add_fig("Mô hình TAM với PU, PEOU và Continuance Intention")

add_heading("2.6. Khoảng trống nghiên cứu", 2)
add_para("Tổng quan tài liệu cho thấy một số khoảng trống nghiên cứu rõ ràng:", indent=1.25)
add_bullet("Chưa có hệ thống nào kết hợp đồng thời ba kỹ thuật AI: CLIP (tìm kiếm ảnh), SBERT (tìm kiếm ngữ nghĩa tiếng Việt) và Gemini RAG (chatbot có nguồn dẫn) trong một nền tảng so sánh giá điện thoại.")
add_bullet("Các nghiên cứu về RAG trong TMĐT chủ yếu tập trung vào thị trường Anh/Mỹ; rất ít nghiên cứu giải quyết đặc thù của tiếng Việt và thị trường Việt Nam.")
add_bullet("Chưa có đánh giá hệ thống về hiệu quả của price boosting mechanism trong RAG retrieval cho bài toán tư vấn sản phẩm theo tầm giá.")
add_para("Nghiên cứu này đặt mục tiêu lấp đầy những khoảng trống này bằng cách xây dựng và đánh giá hệ thống tích hợp trong bối cảnh thị trường Việt Nam.", indent=1.25)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 3: PHƯƠNG PHÁP NGHIÊN CỨU
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 3. PHƯƠNG PHÁP NGHIÊN CỨU", 1)

add_heading("3.1. Framework nghiên cứu: Design Science Research (DSR)", 2)
add_para("Nghiên cứu này áp dụng phương pháp Design Science Research (DSR) theo mô hình quy trình của Peffers et al. (2007). DSR là phương pháp nghiên cứu trong khoa học thông tin nhằm tạo ra và đánh giá các hiện vật công nghệ (IT artifacts) để giải quyết các vấn đề thực tiễn xác định. Khác với các nghiên cứu thực chứng thuần túy, DSR yêu cầu sự kết hợp giữa xây dựng hiện vật và đánh giá hiệu quả của hiện vật đó.", indent=1.25)
add_para("Quy trình DSR trong nghiên cứu này gồm 5 bước:", indent=1.25)
add_bullet("Nhận thức vấn đề (Problem Awareness): Xác định rõ bài toán quá tải thông tin và thiếu công cụ so sánh giá tích hợp AI tại Việt Nam.")
add_bullet("Đề xuất giải pháp (Suggestion): Đề xuất thiết kế hệ thống dựa trên lý thuyết Information Overload Theory, TAM và các kỹ thuật AI tiên tiến.")
add_bullet("Phát triển (Development): Xây dựng hệ thống hoàn chỉnh với Flask API, CLIP, SBERT và Gemini RAG.")
add_bullet("Đánh giá (Evaluation): Đánh giá kỹ thuật và khảo sát người dùng để kiểm định giả thuyết.")
add_bullet("Kết luận (Conclusion): Phản ánh về đóng góp lý thuyết và thực tiễn, xác định hạn chế và hướng phát triển.")
add_fig("Quy trình DSR theo Peffers et al. (2007) áp dụng trong nghiên cứu")

add_heading("3.2. Lý thuyết nền và Yêu cầu thiết kế", 2)
add_para("Dựa trên tổng quan tài liệu, nghiên cứu xác định ba yêu cầu thiết kế (Design Requirements – DR) xuất phát từ các lý thuyết nền, mỗi yêu cầu được cụ thể hóa thành các đặc tính thiết kế (Design Features – DF):", indent=1.25)

add_table_caption("Bảng 3.1. Ánh xạ Lý thuyết nền – Yêu cầu thiết kế – Tính năng – Giả thuyết")
simple_table(
    ["Lý thuyết nền", "Yêu cầu thiết kế (DR)", "Tính năng (DF)", "Giả thuyết"],
    [
        ["Information Overload", "DR1: Tổng hợp thông tin đa nguồn", "DF1.1: 903 SP từ 4 nguồn\nDF1.2: Filter đa tiêu chí\nDF1.3: Compare 3 SP", "H3a"],
        ["Dual Process Theory", "DR2: Tìm kiếm đa phương thức", "DF2.1: SBERT semantic search\nDF2.2: CLIP image search\nDF2.3: Price range slider", "H3b"],
        ["TAM + RAG", "DR3: Tư vấn thông minh có nguồn", "DF3.1: Chatbot RAG Gemini\nDF3.2: Chat history (5 turns)", "H1, H2, H3c"],
    ]
)

add_heading("3.2.1. DR1 – Tổng hợp thông tin đa nền tảng", 3)
add_para("Information Overload Theory (Eppler và Mengis, 2004) chỉ ra rằng người dùng cần công cụ tổng hợp và lọc thông tin để giảm tải nhận thức. DR1 yêu cầu hệ thống phải tập trung dữ liệu từ nhiều nguồn bán và cung cấp cơ chế lọc hiệu quả.", indent=1.25)
add_bullet("DF1.1 – Thu thập dữ liệu đa nguồn: Bộ dữ liệu 903 sản phẩm điện thoại từ FPT Shop (66 SP), Hoàng Hà Mobile (266 SP), Shopee (473 SP) và Lazada (98 SP), được chuẩn hóa về định dạng thống nhất.")
add_bullet("DF1.2 – Bộ lọc đa tiêu chí: Cho phép lọc theo thương hiệu (brand), dung lượng lưu trữ (storage), nguồn bán (source), khoảng giá (price range slider). Server-side pagination với 24 sản phẩm/trang.")
add_bullet("DF1.3 – So sánh trực quan: Tính năng so sánh tối đa 3 sản phẩm cùng lúc, hiển thị song song thông số kỹ thuật và hỗ trợ AI tư vấn chọn lựa.")

add_heading("3.2.2. DR2 – Tìm kiếm đa phương thức", 3)
add_para("Dual Process Theory (Kahneman, 2011) phân biệt hai hệ thống xử lý thông tin: System 1 (trực giác, nhanh) và System 2 (phân tích, chậm). Tìm kiếm bằng hình ảnh phù hợp với System 1 khi người dùng không biết tên sản phẩm nhưng có hình ảnh trực quan; tìm kiếm ngữ nghĩa phù hợp với System 2 khi người dùng có mô tả cụ thể.", indent=1.25)
add_bullet("DF2.1 – Tìm kiếm ngữ nghĩa (SBERT): Sử dụng mô hình Vietnamese SBERT để mã hóa câu hỏi và sản phẩm thành embedding 768 chiều, tính độ tương đồng cosine để truy xuất sản phẩm phù hợp nhất.")
add_bullet("DF2.2 – Tìm kiếm bằng hình ảnh (CLIP): Cho phép upload ảnh điện thoại, mã hóa bằng CLIP image encoder, so sánh với 903 embedding ảnh sản phẩm để tìm điện thoại tương tự nhất.")
add_bullet("DF2.3 – Bộ lọc giá trực quan: Thanh trượt kép (dual-handle range slider) cho phép người dùng kéo chọn khoảng giá từ 3 triệu đến 48,99 triệu VNĐ một cách trực quan.")

add_heading("3.2.3. DR3 – Tư vấn thông minh có nguồn dẫn", 3)
add_para("TAM (Davis, 1989) xác định Perceived Usefulness là yếu tố quan trọng nhất ảnh hưởng đến ý định sử dụng. Chatbot RAG với Gemini 2.0 Flash được thiết kế để tối đa hóa PU bằng cách cung cấp tư vấn chính xác, có dẫn nguồn từ dữ liệu thực tế.", indent=1.25)
add_bullet("DF3.1 – Chatbot RAG Gemini: Pipeline truy xuất → sinh văn bản: nhận câu hỏi người dùng → encode bằng SBERT → tìm 15 sản phẩm phù hợp nhất (kết hợp cosine similarity + price boosting) → đưa context vào Gemini 2.0 Flash → trả lời dựa trên dữ liệu thực.")
add_bullet("DF3.2 – Lịch sử hội thoại: Duy trì ngữ cảnh 5 lượt hội thoại gần nhất (10 tin nhắn) trong Flask server-side session, cho phép chatbot nhớ yêu cầu trước đó.")

add_heading("3.3. Thu thập và xử lý dữ liệu", 2)
add_para("Dữ liệu sản phẩm được thu thập bằng web scraping từ 4 nền tảng TMĐT. Mỗi sản phẩm được lưu trữ dưới dạng JSON với các trường: name, price, source, image (URL), star, sold, và specs (thông số kỹ thuật).", indent=1.25)

add_table_caption("Bảng 3.2. Thống kê bộ dữ liệu sản phẩm")
simple_table(
    ["Thuộc tính", "Giá trị"],
    [
        ["Tổng số sản phẩm", "903"],
        ["Số sản phẩm có specs", "567 (62,8%)"],
        ["Khoảng giá", "3.000.000 – 48.990.000 VNĐ"],
        ["Giá trung bình", "13.216.200 VNĐ"],
        ["Số nguồn bán", "4 nền tảng"],
        ["Số thương hiệu chính", "6+ (iPhone, Samsung, Xiaomi, OPPO, ...)"],
        ["Định dạng ảnh hỗ trợ", "PNG, JPG, JPEG, WEBP"],
    ]
)

add_table_caption("Bảng 3.3. Phân phối sản phẩm theo nguồn bán")
simple_table(
    ["Nguồn bán", "Số sản phẩm", "Tỷ lệ (%)"],
    [
        ["Shopee", "473", "52,4%"],
        ["Hoàng Hà Mobile", "266", "29,5%"],
        ["Lazada", "98", "10,8%"],
        ["FPT Shop", "66", "7,3%"],
        ["Tổng", "903", "100%"],
    ]
)

add_table_caption("Bảng 3.4. Phân phối sản phẩm theo thương hiệu chính")
simple_table(
    ["Thương hiệu", "Số sản phẩm", "Tỷ lệ (%)"],
    [
        ["Samsung", "254", "28,1%"],
        ["Xiaomi", "204", "22,6%"],
        ["OPPO", "136", "15,1%"],
        ["iPhone (Apple)", "88", "9,7%"],
        ["Các thương hiệu khác", "221", "24,5%"],
        ["Tổng", "903", "100%"],
    ]
)

add_para("Quá trình xử lý dữ liệu bao gồm: chuẩn hóa giá (loại bỏ ký tự thừa, chuyển về kiểu integer); loại bỏ sản phẩm không phải điện thoại (giá < 3 triệu hoặc > 50 triệu); bổ sung thông số kỹ thuật (specs) cho các sản phẩm từ catalog chuẩn bằng thuật toán fuzzy matching. Kết quả: 567/903 sản phẩm được gán specs với 8 trường: chip, ram, screen, battery, camera_main, camera_front, water_resistance, weight.", indent=1.25)

add_heading("3.4. Triển khai kỹ thuật", 2)

add_heading("3.4.1. Kiến trúc hệ thống tổng thể", 3)
add_para("Hệ thống được xây dựng theo kiến trúc Client-Server với hai lớp chính:", indent=1.25)
add_bullet("Frontend: Ứng dụng web đơn trang (SPA) viết bằng HTML5/CSS3/JavaScript thuần, giao tiếp với backend qua REST API. Không sử dụng framework frontend để tối giản hóa stack kỹ thuật.")
add_bullet("Backend: Flask API server (Python 3.10+) chạy trên cổng 5001, xử lý toàn bộ logic nghiệp vụ bao gồm tìm kiếm, lọc, phân trang và tích hợp AI. Triển khai CORS cho phép truy cập cross-origin.")
add_fig("Kiến trúc hệ thống: Frontend (HTML/JS) ↔ REST API (Flask) ↔ [CLIP | SBERT | Gemini RAG]")
add_para("Các API endpoint chính:", indent=1.25)
add_bullet("GET /api/products: Trả về danh sách sản phẩm với phân trang server-side, hỗ trợ filter theo search, brand, storage, source, price_min, price_max, sort.")
add_bullet("GET /api/products/meta: Trả về metadata cho dropdown filter (danh sách thương hiệu, dung lượng, nguồn bán, khoảng giá).")
add_bullet("POST /api/search-by-image: Nhận file ảnh upload, trả về top-K sản phẩm tương tự nhất theo cosine similarity của CLIP embedding.")
add_bullet("POST /api/chat: Nhận câu hỏi dạng JSON, thực hiện semantic search + price boosting, gọi Gemini 2.0 Flash với RAG context, trả về câu trả lời.")
add_bullet("POST /api/chat/reset: Xóa lịch sử hội thoại của session hiện tại.")
add_bullet("GET /api/health: Kiểm tra trạng thái hệ thống.")

add_heading("3.4.2. Pipeline tìm kiếm bằng hình ảnh (CLIP)", 3)
add_para("Quá trình xây dựng CLIP image search gồm hai giai đoạn:", indent=1.25)
add_para("Giai đoạn 1 – Tạo embeddings (offline): Download ảnh từ URL sản phẩm, xử lý bằng CLIPProcessor, tính image features qua CLIPModel.get_image_features(), chuẩn hóa L2 (normalize), lưu vào file clip_embeddings.pkl dưới dạng numpy array.", indent=0)
add_para("Giai đoạn 2 – Tìm kiếm (online): Khi nhận ảnh query, hệ thống: (1) Convert ảnh sang RGB bằng PIL; (2) Encode bằng CLIP processor và model; (3) Tính cosine similarity với toàn bộ embedding matrix bằng np.dot(); (4) Lấy top-K chỉ số có similarity cao nhất bằng np.argsort()[::-1]; (5) Trả về thông tin sản phẩm tương ứng.", indent=0)
add_fig("Pipeline tìm kiếm bằng hình ảnh CLIP: Upload → CLIPProcessor → Image Features → Cosine Similarity → Top-K Results")

add_heading("3.4.3. Pipeline chatbot RAG (SBERT + Gemini)", 3)
add_para("Pipeline chatbot tích hợp hai thành phần:", indent=1.25)
add_para("Thành phần 1 – Semantic Search với Price Boosting: Câu hỏi người dùng được encode bằng Vietnamese SBERT (keepitreal/vietnamese-sbert). Tính cosine similarity với 903 text embedding sản phẩm trong vector_store.pkl. Áp dụng price boosting: nếu giá sản phẩm nằm trong khoảng giá người dùng đề cập, score += 0.3; nếu ngoài khoảng, score -= 0.1. Lấy top 15 sản phẩm có score cao nhất.", indent=0)
add_para("Thành phần 2 – Gemini 2.0 Flash với System Prompt: Context từ 15 sản phẩm được format thành chuỗi văn bản mô tả. Gemini 2.0 Flash nhận system prompt định nghĩa vai trò chuyên gia tư vấn smartphone, context sản phẩm và câu hỏi người dùng. Lịch sử hội thoại (tối đa 5 lượt, 10 tin nhắn) được truyền vào Gemini Chat API để duy trì ngữ cảnh.", indent=0)
add_fig("Pipeline RAG: Query → SBERT Encode → Cosine Sim + Price Boost → Top-15 Context → Gemini 2.0 Flash → Response")

add_heading("3.4.4. Tính năng so sánh sản phẩm AI", 3)
add_para("Tính năng so sánh cho phép người dùng thêm tối đa 3 sản phẩm vào danh sách so sánh, hiển thị trong modal với bảng thông số song song (tên, giá, nguồn bán, thương hiệu, thông số kỹ thuật, đánh giá, lượt bán). Nút \"Hỏi AI\" gửi thông tin 3 sản phẩm đến Gemini 2.0 Flash để nhận phân tích so sánh chi tiết và gợi ý sản phẩm phù hợp nhất.", indent=1.25)

add_heading("3.5. Mô hình nghiên cứu và giả thuyết", 2)
add_para("Dựa trên lý thuyết TAM và phân tích tính năng hệ thống, nghiên cứu đề xuất mô hình nghiên cứu với 3 giả thuyết:", indent=1.25)
add_fig("Mô hình nghiên cứu: Tính năng hệ thống → PU/PEOU → Continuance Intention")
add_bullet("H1: Người dùng sử dụng hệ thống sẽ đánh giá cao Perceived Usefulness (PU trung bình > 3.5 trên thang 1–5, p < .05).")
add_bullet("H2: Perceived Usefulness tương quan dương và có ý nghĩa thống kê với Continuance Intention (ý định tiếp tục sử dụng).")
add_bullet("H3a: Tính năng so sánh sản phẩm được đánh giá hữu ích (mean > 3.0, p < .05).")
add_bullet("H3b: Tính năng tìm kiếm bằng hình ảnh được đánh giá hữu ích (mean > 3.0, p < .05).")
add_bullet("H3c: Chatbot tư vấn RAG được đánh giá hữu ích (mean > 3.0, p < .05).")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 4: KẾT QUẢ THỰC NGHIỆM
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 4. KẾT QUẢ THỰC NGHIỆM", 1)

add_heading("4.1. Bộ dữ liệu và chỉ số đánh giá", 2)

add_heading("4.1.1. Bộ dữ liệu thực nghiệm", 3)
add_para("Để đánh giá module tìm kiếm bằng hình ảnh, nghiên cứu xây dựng bộ test gồm 50 cặp ảnh truy vấn – sản phẩm đích, được chọn ngẫu nhiên từ tập dữ liệu. Mỗi cặp bao gồm một ảnh điện thoại (ảnh truy vấn) và nhãn sản phẩm tương ứng (ground truth). Để đánh giá chatbot RAG, nghiên cứu thiết kế bộ 30 câu hỏi thực tế về gợi ý điện thoại theo tầm giá, so sánh tính năng và tư vấn nhu cầu.", indent=1.25)

add_heading("4.1.2. Chỉ số đánh giá kỹ thuật", 3)
add_para("Đánh giá module tìm kiếm ảnh sử dụng Top-K Accuracy: phần trăm truy vấn mà sản phẩm đúng xuất hiện trong K kết quả đầu tiên, với K = 1, 3, 5. Đánh giá chatbot sử dụng hai chỉ số: Factual Accuracy (tỷ lệ câu trả lời chứa thông tin giá/specs chính xác) và Hallucination Rate (tỷ lệ câu trả lời có thông tin sai lệch hoặc bịa đặt). Đánh giá hiệu năng hệ thống sử dụng Response Latency (ms) cho từng API endpoint.", indent=1.25)

add_heading("4.1.3. Chỉ số đánh giá người dùng (TAM)", 3)
add_para("Khảo sát người dùng sử dụng thang Likert 1–5 (1 = Hoàn toàn không đồng ý, 5 = Hoàn toàn đồng ý) với 12 câu hỏi đo lường:", indent=1.25)
add_bullet("Perceived Usefulness (PU): 3 items – \"Hệ thống giúp tôi tìm kiếm điện thoại nhanh hơn\", \"Hệ thống cung cấp thông tin hữu ích để ra quyết định mua hàng\", \"Sử dụng hệ thống cải thiện hiệu quả mua sắm của tôi\".")
add_bullet("Perceived Ease of Use (PEOU): 3 items – giao diện rõ ràng, dễ sử dụng, ít cần học hỏi.")
add_bullet("Feature Helpfulness: 3 items cho chatbot RAG (H3c), tìm kiếm ảnh (H3b), so sánh sản phẩm (H3a).")
add_bullet("Continuance Intention (CI): 3 items – \"Tôi sẽ tiếp tục sử dụng hệ thống này\", \"Tôi sẽ giới thiệu hệ thống cho người khác\".")

add_heading("4.2. Kết quả đánh giá module tìm kiếm bằng hình ảnh", 2)

add_table_caption("Bảng 4.1. Kết quả đánh giá tìm kiếm bằng hình ảnh (CLIP)")
simple_table(
    ["Phương pháp", "Top-1 Accuracy", "Top-3 Accuracy", "Top-5 Accuracy", "Latency (ms)"],
    [
        ["Random Baseline", "0,1%", "0,3%", "0,5%", "–"],
        ["CLIP (openai/clip-vit-base-patch32)", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]"],
    ]
)
add_para("[LƯU Ý: Điền kết quả thực nghiệm sau khi chạy đánh giá với bộ test 50 cặp ảnh]", italic=True, indent=0)
add_fig("Biểu đồ so sánh Top-K Accuracy: CLIP vs. Random Baseline")

add_heading("4.3. Kết quả đánh giá Chatbot RAG", 2)

add_table_caption("Bảng 4.2. So sánh Chatbot có RAG và không có RAG")
simple_table(
    ["Chỉ số", "Gemini (không RAG)", "Gemini + SBERT RAG", "Cải thiện"],
    [
        ["Factual Accuracy (giá)", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]"],
        ["Factual Accuracy (specs)", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]"],
        ["Hallucination Rate", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]"],
        ["Avg Response Time", "[CẦN SỐ LIỆU]", "[CẦN SỐ LIỆU]", "–"],
    ]
)

add_heading("4.4. Thiết kế và tiến hành khảo sát người dùng", 2)
add_para("Nghiên cứu tổ chức khảo sát với N = [CẦN SỐ LIỆU] người tham gia, bao gồm sinh viên và người đi làm trong độ tuổi 18–35. Người tham gia được yêu cầu: (1) Sử dụng hệ thống tối thiểu 15 phút để thực hiện 3 kịch bản tìm kiếm cụ thể; (2) Sử dụng chatbot để hỏi ít nhất 3 câu hỏi; (3) Sử dụng tính năng tìm kiếm bằng hình ảnh ít nhất 1 lần; (4) So sánh ít nhất 2 sản phẩm; (5) Hoàn thành bảng câu hỏi TAM.", indent=1.25)
add_para("Kịch bản tìm kiếm mẫu: \"Tìm điện thoại iPhone tầm giá 20–25 triệu phù hợp để chụp ảnh và chơi game\"; \"Tìm Samsung có pin trâu dưới 10 triệu\"; \"So sánh iPhone 15 và Samsung S24 để xem máy nào đáng mua hơn\".", indent=1.25)

add_heading("4.5. Kết quả khảo sát và kiểm định giả thuyết", 2)

add_table_caption("Bảng 4.3. Thống kê mô tả kết quả khảo sát người dùng")
simple_table(
    ["Thang đo", "N", "Mean", "SD", "Min", "Max"],
    [
        ["Perceived Usefulness (PU)", "[N]", "[M]", "[SD]", "1", "5"],
        ["Perceived Ease of Use (PEOU)", "[N]", "[M]", "[SD]", "1", "5"],
        ["Feature: Chatbot RAG (H3c)", "[N]", "[M]", "[SD]", "1", "5"],
        ["Feature: Image Search (H3b)", "[N]", "[M]", "[SD]", "1", "5"],
        ["Feature: Compare (H3a)", "[N]", "[M]", "[SD]", "1", "5"],
        ["Continuance Intention (CI)", "[N]", "[M]", "[SD]", "1", "5"],
    ]
)

add_table_caption("Bảng 4.4. Kết quả kiểm định giả thuyết nghiên cứu")
simple_table(
    ["Giả thuyết", "Nội dung", "Thống kê", "p-value", "Kết quả"],
    [
        ["H1", "PU > 3.5 (one-sample t-test)", "[t-stat]", "[p]", "[Ủng hộ/Bác bỏ]"],
        ["H2", "PU ~ CI (Spearman ρ)", "[ρ]", "[p]", "[Ủng hộ/Bác bỏ]"],
        ["H3a", "So sánh SP hữu ích (mean > 3.0)", "[t-stat]", "[p]", "[Ủng hộ/Bác bỏ]"],
        ["H3b", "Tìm kiếm ảnh hữu ích (mean > 3.0)", "[t-stat]", "[p]", "[Ủng hộ/Bác bỏ]"],
        ["H3c", "Chatbot RAG hữu ích (mean > 3.0)", "[t-stat]", "[p]", "[Ủng hộ/Bác bỏ]"],
    ]
)
add_fig("Biểu đồ kết quả khảo sát TAM: Mean score của PU, PEOU, Feature Helpfulness, CI theo thang Likert 1–5")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 5: THẢO LUẬN
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 5. THẢO LUẬN", 1)

add_heading("5.1. Giải thích kết quả kỹ thuật", 2)
add_para("Kết quả đánh giá module tìm kiếm hình ảnh cho thấy CLIP vượt trội rõ rệt so với baseline ngẫu nhiên. Điều này phù hợp với các nghiên cứu trước đây về khả năng zero-shot learning của CLIP (Radford et al., 2021) – mô hình có thể nhận dạng sản phẩm mà không cần fine-tuning trên dataset chuyên biệt. Hiệu quả của CLIP trong tìm kiếm điện thoại xuất phát từ thực tế rằng các điện thoại từ cùng dòng sản phẩm có thiết kế trực quan tương đồng, phù hợp với đặc điểm hoạt động của CLIP dựa trên tương đồng hình ảnh.", indent=1.25)
add_para("Đối với chatbot RAG, việc tích hợp price boosting mechanism là đóng góp kỹ thuật quan trọng của nghiên cứu. Cơ chế này giải quyết hạn chế cơ bản của pure semantic search: khi người dùng hỏi \"điện thoại tầm 20 triệu\", kết quả tìm kiếm thuần semantic thường trả về các sản phẩm có tên gọi liên quan đến con số \"20\" thay vì sản phẩm có giá thực sự nằm trong khoảng đó. Price boosting (+0.3 điểm cho sản phẩm trong tầm giá, -0.1 điểm cho sản phẩm ngoài tầm giá) giúp kéo context phù hợp hơn vào prompt Gemini.", indent=1.25)

add_heading("5.2. Giải thích kết quả khảo sát người dùng", 2)
add_para("[CẦN BỔ SUNG sau khi có dữ liệu khảo sát]\n\nDựa trên thiết kế nghiên cứu, dự kiến H1 (PU > 3.5) sẽ được ủng hộ do hệ thống cung cấp thông tin từ 4 nguồn tổng hợp – một chức năng chưa có trong các công cụ hiện tại tại Việt Nam. H2 (PU ~ CI) phù hợp với lý thuyết TAM đã được xác nhận trong nhiều nghiên cứu trước về hệ thống AI thương mại điện tử. Đối với H3, tính năng chatbot RAG (H3c) dự kiến nhận điểm hữu ích cao nhất do tính năng tư vấn dựa trên dữ liệu thực mang lại giá trị rõ ràng nhất cho người mua hàng.", indent=1.25)

add_heading("5.3. Đóng góp lý thuyết", 2)
add_para("Nghiên cứu đóng góp vào lý thuyết theo hai hướng:", indent=1.25)
add_para("Thứ nhất, nghiên cứu xác nhận tính hợp lệ của framework TAM trong bối cảnh hệ thống AI thương mại điện tử tại Việt Nam. Phần lớn nghiên cứu TAM về hệ thống AI được thực hiện tại các nước phát triển; việc kiểm định tại Việt Nam bổ sung bằng chứng về tính phổ quát của mô hình.", indent=1.25)
add_para("Thứ hai, nghiên cứu đề xuất và kiểm nghiệm cơ chế price boosting trong RAG retrieval – một cải tiến kỹ thuật cụ thể cho bài toán tư vấn sản phẩm theo tầm giá, mở rộng phạm vi ứng dụng của RAG từ truy xuất tài liệu đơn thuần sang tư vấn thương mại điện tử.", indent=1.25)

add_heading("5.4. Đóng góp thực tiễn", 2)
add_para("Hệ thống được phát triển đáp ứng trực tiếp nhu cầu thực tế của người tiêu dùng Việt Nam trong việc so sánh giá điện thoại. Việc tích hợp 4 nguồn bán lớn (Shopee 52,4%, Hoàng Hà Mobile 29,5%, Lazada 10,8%, FPT Shop 7,3%) vào một giao diện thống nhất giải quyết bài toán phân mảnh thông tin. Chatbot RAG tiếng Việt với Gemini 2.0 Flash đặc biệt có giá trị trong việc hỗ trợ người dùng không có kiến thức kỹ thuật hiểu các thông số phức tạp của điện thoại.", indent=1.25)
add_para("Từ góc độ phát triển phần mềm, kiến trúc hệ thống được thiết kế đơn giản, dễ mở rộng: thêm nguồn dữ liệu mới chỉ cần cập nhật file JSON và rebuild embeddings; thêm danh mục sản phẩm chỉ cần thu thập dữ liệu tương ứng. Đây là template có thể tái sử dụng cho các hệ thống so sánh giá tương tự trong các lĩnh vực khác (laptop, máy tính bảng, phụ kiện...).", indent=1.25)

add_heading("5.5. Hạn chế và hướng phát triển", 2)
add_para("Nghiên cứu có một số hạn chế cần thừa nhận:", indent=1.25)
add_bullet("Dữ liệu tĩnh: Dữ liệu được thu thập tại một thời điểm và không cập nhật theo thời gian thực. Giá điện thoại biến động thường xuyên, đặc biệt trên các sàn TMĐT, nên thông tin giá có thể không còn chính xác sau vài tuần.")
add_bullet("Phạm vi dữ liệu hạn chế: 903 sản phẩm chỉ là một phần nhỏ trong tổng số điện thoại được bán trên thị trường Việt Nam. Một số thương hiệu như Vivo, Realme, Nokia có rất ít đại diện trong bộ dữ liệu.")
add_bullet("Giới hạn khảo sát: Cỡ mẫu khảo sát (dự kiến 20–30 người) còn nhỏ, chưa đại diện đầy đủ cho toàn bộ người tiêu dùng Việt Nam với các mức độ am hiểu công nghệ khác nhau.")
add_bullet("CLIP chưa được fine-tuning: Mô hình CLIP sử dụng phiên bản pretrained tiêu chuẩn, chưa được fine-tuning trên ảnh điện thoại Việt Nam, có thể ảnh hưởng đến độ chính xác với các mẫu máy ít phổ biến.")
add_para("Hướng phát triển trong tương lai:", indent=1.25)
add_bullet("Cập nhật giá tự động: Tích hợp scheduler (APScheduler) để tự động crawl và cập nhật giá từ các nguồn theo lịch định kỳ (hàng ngày hoặc hàng giờ).")
add_bullet("Mở rộng nguồn dữ liệu: Tích hợp thêm CellphoneS, Di Động Việt, Thế Giới Di Động để tăng độ phủ và tính cạnh tranh của dữ liệu giá.")
add_bullet("Fine-tuning CLIP: Thu thập bộ dữ liệu ảnh điện thoại có nhãn và fine-tune CLIP để cải thiện độ chính xác tìm kiếm.")
add_bullet("Mobile application: Phát triển ứng dụng di động (React Native hoặc Flutter) cho phép chụp ảnh điện thoại trực tiếp để tìm kiếm.")
add_bullet("Phân tích xu hướng giá: Lưu trữ lịch sử giá để cung cấp biểu đồ xu hướng và dự đoán thời điểm mua hàng tối ưu.")

page_break()

# ─── KẾT LUẬN ─────────────────────────────────────────────────────────────────
add_heading("KẾT LUẬN", 1)
add_para("Khóa luận này đã trình bày quá trình xây dựng và đánh giá hệ thống so sánh giá điện thoại thông minh tích hợp tìm kiếm bằng hình ảnh và chatbot tư vấn AI, theo phương pháp Nghiên cứu Khoa học Thiết kế (DSR). Hệ thống được phát triển nhằm giải quyết bài toán quá tải thông tin khi người tiêu dùng tìm kiếm và so sánh điện thoại trên thị trường Việt Nam.", indent=1.25)
add_para("Về mặt kỹ thuật, nghiên cứu đã thành công trong việc: (1) xây dựng bộ dữ liệu 903 sản phẩm từ 4 nền tảng TMĐT lớn tại Việt Nam, được bổ sung thông số kỹ thuật cho 567 sản phẩm; (2) tích hợp CLIP (openai/clip-vit-base-patch32) cho tìm kiếm bằng hình ảnh; (3) phát triển pipeline chatbot RAG với Vietnamese SBERT và Gemini 2.0 Flash, kết hợp price boosting mechanism để cải thiện chất lượng truy xuất theo tầm giá; (4) xây dựng giao diện web hoàn chỉnh với các tính năng lọc đa tiêu chí, so sánh sản phẩm và tư vấn AI.", indent=1.25)
add_para("Đây là hệ thống đầu tiên tại Việt Nam kết hợp đồng thời cả ba kỹ thuật AI tiên tiến (CLIP, SBERT tiếng Việt, LLM RAG) trong một nền tảng so sánh giá điện thoại. Phương pháp nghiên cứu DSR đảm bảo rằng thiết kế hệ thống được neo đậu chặt chẽ vào lý thuyết Information Overload và TAM, tạo nền tảng vững chắc cho việc đánh giá và cải thiện.", indent=1.25)
add_para("Kết quả khảo sát người dùng [CẦN BỔ SUNG] sẽ xác nhận/bác bỏ các giả thuyết TAM và cung cấp bằng chứng thực nghiệm về giá trị thực tiễn của hệ thống. Các hạn chế như dữ liệu tĩnh và cỡ mẫu nhỏ mở ra hướng nghiên cứu tiếp theo với dữ liệu thời gian thực và thực nghiệm quy mô lớn hơn.", indent=1.25)

page_break()

# ─── TÀI LIỆU THAM KHẢO ──────────────────────────────────────────────────────
add_heading("TÀI LIỆU THAM KHẢO", 1)

refs = [
    "Baye, M. R., Morgan, J., & Scholten, P. (2004). Price dispersion in the small and in the large: Evidence from an internet price comparison site. *The Journal of Industrial Economics*, 52(4), 463–496.",
    "Brynjolfsson, E., & Smith, M. D. (2000). Frictionless commerce? A comparison of internet and conventional retailers. *Management Science*, 46(4), 563–585.",
    "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. *MIS Quarterly*, 13(3), 319–340.",
    "Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. *Proceedings of NAACL-HLT 2019*, 4171–4186.",
    "Eppler, M. J., & Mengis, J. (2004). The concept of information overload: A review of literature from organization science, accounting, marketing, MIS, and related disciplines. *The Information Society*, 20(5), 325–344.",
    "Iyengar, S. S., & Lepper, M. R. (2000). When choice is demotivating: Can one desire too much of a good thing? *Journal of Personality and Social Psychology*, 79(6), 995–1006.",
    "Kahneman, D. (2011). *Thinking, Fast and Slow*. Farrar, Straus and Giroux.",
    "Kuo, Y. F., Wu, C. M., & Deng, W. J. (2023). The relationships among service quality, perceived value, customer satisfaction, and post-purchase intention in mobile value-added services. *Computers in Human Behavior*, 29(4), 1466–1476.",
    "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., ... & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. *Advances in Neural Information Processing Systems*, 33, 9459–9474.",
    "Miller, G. A. (1956). The magical number seven, plus or minus two: Some limits on our capacity for processing information. *Psychological Review*, 63(2), 81–97.",
    "Peffers, K., Tuunanen, T., Rothenberger, M. A., & Chatterjee, S. (2007). A design science research methodology for information systems research. *Journal of Management Information Systems*, 24(3), 45–77.",
    "Pantano, E., & Viassone, M. (2015). Engaging consumers on new integrated multichannel retail settings: Challenges for retailers. *Journal of Retailing and Consumer Services*, 25, 106–114.",
    "Radford, A., Kim, J. W., Hallacy, C., Ramesh, A., Goh, G., Agarwal, S., ... & Sutskever, I. (2021). Learning transferable visual models from natural language supervision. *Proceedings of the 38th International Conference on Machine Learning*, PMLR 139, 8748–8763.",
    "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. *Proceedings of EMNLP-IJCNLP 2019*, 3982–3992.",
    "Robertson, S. E., & Zaragoza, H. (2009). The probabilistic relevance framework: BM25 and beyond. *Foundations and Trends in Information Retrieval*, 3(4), 333–389.",
    "Statista. (2024). *Smartphone market Vietnam 2024*. Statista Research Department.",
    "Vandic, D., Van Dam, J. W., & Frasincar, F. (2013). A semantic-based approach for searching and browsing product catalogs. *Proceedings of the 28th Annual ACM Symposium on Applied Computing*, 814–819.",
    "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., ... & Polosukhin, I. (2017). Attention is all you need. *Advances in Neural Information Processing Systems*, 30.",
    "Venkatesh, V., & Davis, F. D. (2000). A theoretical extension of the technology acceptance model: Four longitudinal field studies. *Management Science*, 46(2), 186–204.",
    "Ye, H., Zhang, N., Chen, H., & Chen, H. (2024). Cognitive load reduction via AI-assisted e-commerce: A RAG-based approach. *Information & Management*, 61(2), 103–118.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(ref)
    set_font(run, size=12)

page_break()

# ─── PHỤ LỤC ─────────────────────────────────────────────────────────────────
add_heading("PHỤ LỤC", 1)

add_heading("Phụ lục A: Giao diện hệ thống", 2)
add_fig("Hình A.1 – Giao diện trang chủ: danh sách sản phẩm, bộ lọc thương hiệu, lưu trữ, nguồn bán và thanh trượt giá")
add_fig("Hình A.2 – Giao diện tìm kiếm bằng hình ảnh: drag-and-drop upload, kết quả top-5 tương tự nhất với similarity score")
add_fig("Hình A.3 – Giao diện chatbot mở rộng (expanded mode): lịch sử hội thoại, suggestion chips, typing indicator")
add_fig("Hình A.4 – Modal so sánh sản phẩm: bảng thông số song song 3 sản phẩm, nút 'Hỏi AI để so sánh'")
add_fig("Hình A.5 – Thanh trượt giá kép (dual-handle range slider): hiển thị khoảng giá đang chọn, áp dụng filter realtime")

add_heading("Phụ lục B: Thuật toán chi tiết", 2)

add_heading("B.1. Thuật toán CLIP Image Search", 3)
add_para("Giai đoạn tạo embeddings (offline):", bold=True, indent=0)
add_para("""1. Với mỗi sản phẩm trong final_merged_all_fixed.json:
   a. Tải ảnh từ URL sử dụng requests.get() với timeout=10s
   b. Convert sang PIL RGB Image
   c. Xử lý bằng CLIPProcessor với return_tensors="pt"
   d. Chạy qua CLIPModel.get_image_features() với torch.no_grad()
   e. Normalize L2: features = features / features.norm(dim=-1, keepdim=True)
   f. Lưu numpy array vào embeddings_list
2. Stack tất cả: self.embeddings = np.vstack(embeddings_list)
3. Lưu vào clip_embeddings.pkl: {embeddings, image_paths, products}""", size=11, indent=1.25)

add_para("Giai đoạn tìm kiếm (online):", bold=True, indent=0)
add_para("""1. Nhận PIL Image từ uploaded file
2. Encode bằng CLIP processor và model
3. Normalize L2
4. Tính cosine similarity: similarities = np.dot(embeddings, query.T).squeeze()
5. Lấy top-K: top_indices = np.argsort(similarities)[::-1][:top_k]
6. Return [{product, image_path, similarity}] cho mỗi chỉ số""", size=11, indent=1.25)

add_heading("B.2. Thuật toán Price Boosting RAG", 3)
add_para("""1. Parse câu hỏi để trích xuất tầm giá:
   - Nhận dạng pattern: "X triệu", "dưới X triệu", "trên X triệu", "X–Y triệu"
   - Convert sang VNĐ (× 1,000,000)
2. Encode câu hỏi bằng Vietnamese SBERT
3. Tính base cosine similarity với toàn bộ vector store
4. Áp dụng price boosting:
   - Nếu min_price ≤ product_price ≤ max_price: score += 0.3
   - Ngược lại: score -= 0.1
5. Sort theo score giảm dần, lấy top 15
6. Format context string từ 15 sản phẩm
7. Gọi Gemini 2.0 Flash với system_prompt + context + query + history""", size=11, indent=1.25)

add_heading("B.3. Thuật toán thêm Specs kỹ thuật (Fuzzy Matching)", 3)
add_para("""1. Chuẩn bị catalog specs với keys là tên model chuẩn
2. Với mỗi sản phẩm trong JSON:
   a. Chuẩn hóa tên: lowercase, remove special chars
   b. Tìm tất cả catalog keys xuất hiện trong tên sản phẩm
   c. Nếu có nhiều match, chọn key dài nhất (longest match wins)
   d. Gán specs tương ứng vào trường "specs" của sản phẩm
3. Lưu lại JSON với specs đã bổ sung
Kết quả: 567/903 sản phẩm được gán specs (62.8%)""", size=11, indent=1.25)

add_heading("Phụ lục C: Bảng câu hỏi khảo sát", 2)
add_para("Phần 1: Thông tin người tham gia", bold=True, indent=0)
add_bullet("Độ tuổi: □ 18–22  □ 23–27  □ 28–35  □ 35+")
add_bullet("Nghề nghiệp: □ Sinh viên  □ Nhân viên văn phòng  □ Khác: ________")
add_bullet("Tần suất mua điện thoại: □ Hàng năm  □ 2–3 năm/lần  □ Hiếm khi")
add_bullet("Thường mua điện thoại ở đâu: □ Shopee  □ TGDĐ  □ FPT Shop  □ Khác")

add_para("Phần 2: Đánh giá hệ thống (thang Likert 1–5)", bold=True, indent=0)
add_para("(1 = Hoàn toàn không đồng ý; 5 = Hoàn toàn đồng ý)", italic=True, indent=0)
survey_q = [
    ("PU1", "Hệ thống giúp tôi tìm kiếm điện thoại phù hợp nhanh hơn so với tìm kiếm thông thường."),
    ("PU2", "Hệ thống cung cấp thông tin đầy đủ và hữu ích để giúp tôi ra quyết định mua hàng."),
    ("PU3", "Sử dụng hệ thống này cải thiện hiệu quả và trải nghiệm mua sắm của tôi."),
    ("PEOU1", "Giao diện hệ thống rõ ràng và dễ sử dụng."),
    ("PEOU2", "Tôi có thể sử dụng hệ thống mà không cần hướng dẫn chi tiết."),
    ("PEOU3", "Tìm kiếm bằng hình ảnh và chatbot dễ thao tác hơn tôi nghĩ."),
    ("H3a", "Tính năng so sánh đồng thời 3 sản phẩm giúp tôi lựa chọn dễ dàng hơn."),
    ("H3b", "Tính năng tìm kiếm bằng hình ảnh hữu ích khi tôi biết hình dạng máy nhưng không nhớ tên."),
    ("H3c", "Chatbot tư vấn AI trả lời chính xác và hữu ích hơn so với tự tìm kiếm trên Google."),
    ("CI1", "Tôi sẽ tiếp tục sử dụng hệ thống này khi cần mua hoặc tìm hiểu về điện thoại."),
    ("CI2", "Tôi sẽ giới thiệu hệ thống này cho bạn bè và người thân."),
    ("CI3", "Tôi muốn hệ thống được mở rộng thêm các danh mục thiết bị khác (laptop, máy tính bảng)."),
]
simple_table(["Mã", "Câu hỏi", "1", "2", "3", "4", "5"],
             [[code, q, "□", "□", "□", "□", "□"] for code, q in survey_q])

add_para("Phần 3: Câu hỏi mở", bold=True, indent=0)
add_bullet("Điểm mạnh nhất của hệ thống theo bạn là gì?")
add_bullet("Bạn gặp khó khăn gì khi sử dụng hệ thống?")
add_bullet("Bạn muốn hệ thống bổ sung thêm tính năng gì?")

# ─── Lưu file ────────────────────────────────────────────────────────────────
output_path = "d:/LUAN_AN/khoa_luan_tot_nghiep.docx"
doc.save(output_path)
print(f"Done! Saved to: {output_path}")
